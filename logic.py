import os
import subprocess
import tempfile
import io
import re
from pathlib import Path
import soundfile as sf
import numpy as np
import torch
import docx
from transformers import AutoProcessor, AutoModelForCTC
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

folder_id = os.getenv('YANDEX_FOLDER_ID')
client = OpenAI(
    api_key=os.getenv('YANDEX_API_KEY'),
    base_url="https://ai.api.cloud.yandex.net/v1",
    project=folder_id
)
llm_model_name = f"gpt://{folder_id}/aliceai-llm"

MODEL_ID = "AigizK/wav2vec2-large-mms-1b-tatar"
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

chunk_seconds = 20
concat_by_text = True

processor = AutoProcessor.from_pretrained(MODEL_ID)
model = AutoModelForCTC.from_pretrained(MODEL_ID)
model.to(DEVICE)
model.eval()

def convert_to_wav_16k(input_path: str, output_path: str, apply_limit: bool = True):
    cmd = ["ffmpeg", "-y", "-i", input_path]
    if apply_limit:
        cmd.extend(["-t", "180"])
    cmd.extend(["-ar", "16000", "-ac", "1", "-sample_fmt", "s16", output_path])
    
    try:
        proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if proc.returncode != 0:
            err_msg = proc.stderr.decode('utf-8', errors='ignore')
            raise RuntimeError(f"FFmpeg ошибка: {err_msg}")
    except FileNotFoundError:
        raise RuntimeError("Утилита FFmpeg не найдена в системе. Если ты на MacOS, выполни в терминале: brew install ffmpeg")
    
    return output_path

def wav_recognition_with_chunk(wav_path: str) -> str:
    audio, sr = sf.read(wav_path, dtype="float32")
    if len(audio.shape) > 1:
        audio = np.mean(audio, axis=1)

    samples_per_chunk = int(chunk_seconds * sr)
    total_samples = len(audio)
    if total_samples == 0:
        return ""

    logits_list = []
    chunk_texts = []

    model.eval()
    with torch.no_grad():
        for start in range(0, total_samples, samples_per_chunk):
            end = min(start + samples_per_chunk, total_samples)
            chunk = audio[start:end]

            inputs = processor(chunk, sampling_rate=sr, return_tensors="pt", padding=True)
            inputs = {k: v.to(DEVICE) for k, v in inputs.items()}

            outputs = model(**inputs).logits
            logits_np = outputs[0].cpu().numpy()

            if concat_by_text:
                chunk_texts.append(processor.decode(logits_np)['text'])
            else:
                logits_list.append(logits_np)

    if concat_by_text:
        return " ".join(t for t in chunk_texts if t)

    if not logits_list:
        return ""

    all_logits = np.concatenate(logits_list, axis=0)
    return processor.decode(all_logits)['text']

class Interaction_with_LLM:
    def __init__(self, model_name=llm_model_name):
        self.messages = [
            {
                "role": "system",
                "content": "Представь, что ты ассистент, помогающий людям изучать татарский язык."
            }
        ]
        self.model = model_name

    def create_refined(self, input_text):
        self.messages.append({
            "role": "user",
            "content": f"Полученный текст сырой. Нужно сделать так, чтобы он был чистым, с исправленными грамматическими и орфографическими ошибками и корректной пунктуацией. Реконструкция и вольности при восстановлении смысла допустимы для логичности и связности текста. В ответ давай только получившийся текст на татарском, без лишних объяснений. Полученный сырой текст на татарском: {input_text}"
        })
        response = client.chat.completions.create(model=self.model, messages=self.messages)
        result = response.choices[0].message.content
        self.messages.append({"role": "assistant", "content": result})
        return result

    def create_i_frank_lesson(self):
        self.messages.append({
            "role": "user",
            "content": "Возьми восстановленный текст, представленный тобой ранее, переделай его по методу Ильи Франка для изучения иностранных языков (в данном случае - Татарский). Метод Ильи Франка: цельный текст с переводом после каждой смысловой части. Также обязательно нужны пояснения строго про 1-2 татарских слова после перевода смысловой части. Например, по смысловым фрагментам внутри предложения. Реконструкция и вольный перевод допустим для логичности и связности перевода. Пояснения должны быть в одном блоке с фрагментом перевода после каждой смысловой части. Части с переводом и пояснениями должны быть на русском языке в скобках. Иногда пояснения слов можно пропускать, если смысловая часть простая. Ни в коем случае ничего не убирай из распознанной речи. В ответ давай только получившийся текст, не задавай вопросы для уточнения контекста."
        })
        response = client.chat.completions.create(model=self.model, messages=self.messages)
        result = response.choices[0].message.content
        self.messages.append({"role": "assistant", "content": result})
        return result

    def create_collocations_lesson(self):
        self.messages.append({
            "role": "user",
            "content": "Твоя задача: проанализировать восстановленный текст на татарском языке, который ты сгенерировал ранее, СНАЧАЛА выделить его основную тему, а потом — ключевые устойчивые выражения (collocations), связанные с этой темой. ШАГ 1. ОПРЕДЕЛИ ТЕМУ ТЕКСТА: Кратко сформулируй, о чём этот текст. Назови это блоком: 'Главная тема текста'. ШАГ 2. ВЫБЕРИ КЛЮЧЕВЫЕ COLLOCATIONS: Выбери из текста от 7 до 10 ключевых устойчивых выражений. Ответ должен быть на русском языке (кроме самих татарских фраз). Формат вывода — строго Markdown. Никаких вводных слов. Сначала всегда идёт блок про главную тему. Затем список collocations. Никакого JSON."
        })
        response = client.chat.completions.create(model=self.model, messages=self.messages)
        result = response.choices[0].message.content
        self.messages.append({"role": "assistant", "content": result})
        return result

    def create_matching_lesson(self):
        self.messages.append({
            "role": "user",
            "content": "Создать упражнение на сопоставление, опираясь на список выражений, который ты составил в предыдущем сообщении. 1. Взять ВСЕ выражения. 2. Составь два списка: Левый столбец: Татарские фразы (1., 2., 3. ...). Правый столбец: Их переводы на русский (A., B., C. ...). Переводы должны быть ПЕРЕМЕШАНЫ. Формат вывода: Заголовок: 'Закрепим: Найди пары'. Инструкция: 'Соедините цифру и букву. Напишите ответ в чат (например: 1A 2B).' Блок упражнения. Разделитель: '|||'. Блок для проверки (JSON): Объект, где ключ — цифра, значение — буква. Никаких вводных слов. JSON должен быть валидным текстом."
        })
        response = client.chat.completions.create(model=self.model, messages=self.messages)
        result = response.choices[0].message.content
        self.messages.append({"role": "assistant", "content": result})
        return result

    def create_gap_fill_lesson(self):
        self.messages.append({
            "role": "user",
            "content": "Создать тест, где нужно вставить пропущенные фразы в предложения. Используй те же коллокации. Формат вывода: Заголовок: 'Вставьте пропущенные фразы'. Инструкция: 'Выберите подходящее выражение, опираясь на смысл предложения.' Список вопросов с вариантами a, b, c. Разделитель: '|||'. Блок для проверки (JSON)."
        })
        response = client.chat.completions.create(model=self.model, messages=self.messages)
        result = response.choices[0].message.content
        self.messages.append({"role": "assistant", "content": result})
        return result

    def create_logical_completion_lesson(self):
        self.messages.append({
            "role": "user",
            "content": "Создать тест на логическое завершение предложений, используя ВСЕ collocations. Формат вывода: Заголовок: 'Логическое продолжение'. Инструкция: 'Прочитайте начало фразы и выберите наиболее подходящее продолжение.' Список вопросов. Начало предложения строго без перевода. Разделитель: '|||'. Блок для проверки (JSON)."
        })
        response = client.chat.completions.create(model=self.model, messages=self.messages)
        result = response.choices[0].message.content
        self.messages.append({"role": "assistant", "content": result})
        return result

    def create_grammar_digest_lesson(self):
        self.messages.append({
            "role": "user",
            "content": "Создать мини-урок по грамматике. Выбери 3-4 грамматических явления. Формат вывода: Заголовок: 'Как устроен этот текст: Главные кирпичики' (выведи 1 раз в начале). Разделитель `===` (строго три знака равно на отдельной строке). Блок 1 (Теория + Практика). Разделитель `===`. Блок 2 (Теория + Практика). Разделитель `===`. Блок 3 (Теория + Практика). Разделитель: '|||'. Блок для проверки (JSON)."
        })
        response = client.chat.completions.create(model=self.model, messages=self.messages)
        result = response.choices[0].message.content
        self.messages.append({"role": "assistant", "content": result})
        return result

    def setup_avatar_context(self, text: str, collocations: str):
        prompt = f"""Ты — интерактивный репетитор татарского языка (ИИ-аватар). 
Твоя цель: активно развивать диалог с учеником и помогать ему практиковать язык.
Опирайся на этот текст урока: {text}
И эти ключевые слова: {collocations}

Твои строгие правила:
1. Проявляй инициативу: задавай вопросы по теме текста, предлагай разыграть сценку.
2. Мягко исправляй ошибки ученика.
3. ВСЕГДА пиши свои реплики на татарском языке и ОБЯЗАТЕЛЬНО давай перевод на русский в скобках.
4. ТВОЯ ПЕРВАЯ РЕПЛИКА должна быть приветствием и сразу содержать вопрос к ученику по тексту урока (например, понравился ли рецепт, как прошел день и т.д.)."""
        self.messages.append({
            "role": "system",
            "content": prompt
        })

    def generate_first_message(self) -> str:
        self.messages.append({"role": "user", "content": "Начни наш диалог. Поздоровайся и задай мне первый вопрос по тексту."})
        response = client.chat.completions.create(model=self.model, messages=self.messages)
        reply = response.choices[0].message.content
        self.messages.append({"role": "assistant", "content": reply})
        return reply

    def chat_with_avatar(self, user_message: str) -> str:
        self.messages.append({"role": "user", "content": user_message})
        response = client.chat.completions.create(model=self.model, messages=self.messages)
        reply = response.choices[0].message.content
        self.messages.append({"role": "assistant", "content": reply})
        return reply

def generate_word_document(lesson_data: dict) -> bytes:
    doc = docx.Document()
    doc.add_heading('На родном. Татарча', 0)
    
    sections = {
        'refined_text': 'Текст урока',
        'frank_text': 'Метод Ильи Франка',
        'collocations': 'Словарь',
        'matching': 'Сопоставление',
        'gap_fill': 'Заполнение пропусков',
        'logical': 'Логическое продолжение',
        'grammar': 'Грамматика'
    }
    
    for key, title in sections.items():
        if key in lesson_data and lesson_data[key]:
            doc.add_heading(title, level=1)
            clean_text = lesson_data[key].split('|||')[0]
            clean_text = clean_text.replace('===', '\n\n')
            clean_text = clean_text.replace('**', '').replace('*', '').replace('__', '').replace('_', '')
            doc.add_paragraph(clean_text)
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()